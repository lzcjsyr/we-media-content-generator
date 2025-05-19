#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import re
import json
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# 从主模块导入函数
from content_generator import (
    read_input_file, 
    generate_content_with_title
)

# 从content_generator导入模型常量
from content_generator import (
    OPENROUTER_GEMINI_MODEL,
    OPENROUTER_CLAUDE_MODEL,
    OPENROUTER_DS_R1_MODEL
)

def create_comparison_report(results, output_dir):
    """创建包含所有模型生成内容的比较报告"""
    doc = Document()
    
    def set_font(element, size=None):
        """设置元素的字体"""
        if not hasattr(element, 'runs'):
            return
            
        for run in element.runs:
            run.font.name = 'Times New Roman'
            if not run._element.rPr:
                run._element.get_or_add_rPr()
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if size:
                run.font.size = size

    def add_section(doc, title, level=1):
        """添加标题部分"""
        heading = doc.add_heading(title, level=level)
        set_font(heading, size=Pt(18) if level == 0 else None)
        return heading

    def create_table(doc, headers, rows, style='Table Grid'):
        """创建表格并填充数据"""
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = style
        
        # 设置表头
        for i, header in enumerate(headers):
            table.cell(0, i).text = str(header)
        
        # 填充数据
        for i, row in enumerate(rows, 1):
            for j, cell in enumerate(row):
                table.cell(i, j).text = str(cell)
        return table

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)

    # 添加文档标题
    add_section(doc, '模型比较报告', level=0)
    doc.add_paragraph('本报告对比了不同模型生成的内容效果，包括标题和段落内容。')
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%m-%d %H:%M")}\n')

    # 1. 标题比较部分
    add_section(doc, '一、标题比较')
    
    # 准备标题数据
    models = ['Gemini', 'Claude', 'DeepSeek']
    title_data = []
    for i in range(1, 4):
        row = [f'标题 {i}']
        for model in [m.lower() for m in models]:
            content = results.get(model, {})
            if isinstance(content, dict):
                row.append(str(content.get(f'title_{i}', '未生成')))
            else:
                row.append(str(content) if i == 1 else 'N/A')
        title_data.append(row)
    
    # 创建标题比较表格
    create_table(doc, ['标题方案'] + models, title_data)
    doc.add_paragraph()

    # 2. 段落内容比较
    add_section(doc, '二、正文内容比较')
    
    # 获取所有段落的段落数据
    paragraphs_data = {
        model.lower(): results.get(model.lower(), {}).get('paragraphs', []) 
        if isinstance(results.get(model.lower()), dict) 
        else [] 
        for model in models
    }
    max_paragraphs = max(len(paragraphs) for paragraphs in paragraphs_data.values())

    # 遍历所有段落
    for i in range(max_paragraphs):
        add_section(doc, f'段落 {i+1}', level=2)
        
        # 小标题比较
        add_section(doc, '小标题比较:', level=3)
        subtitles = [
            paragraphs[i].get('subtitle', '无小标题') if i < len(paragraphs) else '无内容'
            for paragraphs in paragraphs_data.values()
        ]
        create_table(doc, models, [subtitles])
        doc.add_paragraph()

        # 段落内容比较
        add_section(doc, '段落内容比较:', level=3)
        for model in models:
            add_section(doc, f'{model}:', level=4)
            content = '无内容'
            if i < len(paragraphs_data[model.lower()]):
                content = paragraphs_data[model.lower()][i].get('content', '无内容')
            doc.add_paragraph(content)
        doc.add_paragraph()

    # 保存文档
    doc_path = os.path.join(output_dir, "model_comparison.docx")
    doc.save(doc_path)
    print(f"模型比较报告已保存: {doc_path}")
    return doc_path

def main(reqs="", temperature=0.7):
    # 验证temperature参数
    if not (0 <= temperature <= 1):
        raise ValueError("temperature参数必须在0到1之间")
    # 读取输入
    content = read_input_file("input.txt")
    if not content:
        print("无法读取输入文件 input.txt，程序终止")
        return
    
    # 创建输出目录
    timestamp = datetime.now().strftime("%m%d_%H%M")
    output_dir = f"comparison_{timestamp}"
    os.makedirs(output_dir, exist_ok=True)
    
    # 定义模型和友好名称的映射
    models = {
        "gemini": OPENROUTER_GEMINI_MODEL,
        "claude": OPENROUTER_CLAUDE_MODEL,
        "deepseek": OPENROUTER_DS_R1_MODEL
    }
    
    # 存储各模型的生成结果
    results = {}
    
    # 使用三个不同模型生成内容
    for model_name, model_id in models.items():
        print(f"\n正在使用 {model_name} 模型生成内容...")
        
        # 为每个模型创建单独的子目录
        model_dir = os.path.join(output_dir, model_name)
        os.makedirs(model_dir, exist_ok=True)
        
        # 生成内容
        try:
            result = generate_content_with_title(content, output_dir=model_dir, model=model_id, reqs=reqs, temperature=temperature)
            results[model_name] = result
            
            # 保存JSON结果到子目录
            json_path = os.path.join(model_dir, "content.json")
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
                
            print(f"{model_name} 模型内容已保存到 {json_path}")
            
        except Exception as e:
            print(f"{model_name} 模型生成失败: {e}")
            results[model_name] = {"title_1": f"{model_name} 生成失败", "paragraphs": []}
    
    # 创建比较报告
    if results:
        report_path = create_comparison_report(results, output_dir)
        
        # 打印结果概要
        print("\n模型比较完成!")
        print(f"比较结果已保存到目录: {output_dir}")
        print(f"比较报告: {report_path}")
        
        # 为每个模型打印标题示例
        for model_name in models.keys():
            if model_name in results:
                title = results[model_name].get("title_1", "未生成")
                print(f"{model_name} 标题示例: {title}")
    else:
        print("所有模型生成失败，无法创建比较报告")

if __name__ == "__main__":

    """
        使用额外要求:
        - 在reqs参数中传入字符串，例如：
          main(reqs="文章风格要正式，使用专业术语")
        
        控制生成随机性:
        - temperature: 控制生成文本的随机性，范围0-1，值越大随机性越强，默认0.7
          例如: main(temperature=0.5)  # 更确定性的输出
               main(temperature=0.9)  # 更有创意的输出
    """
    
    main(reqs="文章风格生动活泼，像好友聊天", temperature=0.7)
