import os
import json
import base64
import time
import re
import sys
import requests
from datetime import datetime
from dotenv import load_dotenv
from openai import OpenAI
import json_repair
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from prompts import (
    SYSTEM_PROMPT, IMAGE_STYLE_PROMPT, IMAGE_SIZE, IMAGE_QUALITY, IMAGE_MODERATION, IMAGE_BACKGROUND
)

# 加载环境变量
load_dotenv()

# 环境变量配置
required_vars = [
    "OPENROUTER_API_KEY",
    "OPENROUTER_BASE_URL", 
    "OPENROUTER_Claude_3.7_Sonnet",
    "OPENROUTER_Gemini_2.5_Pro",
    "OPENROUTER_DeepSeek_R1",
    "AIHUBMIX_API_KEY",
    "AIHUBMIX_BASE_URL",
    "AIHUBMIX_IMAGE_GENERATION_MODEL"
]

# 检查必要环境变量
missing_vars = [var for var in required_vars if not os.getenv(var)]
if missing_vars:
    print(f"错误: 缺少环境变量: {', '.join(missing_vars)}")
    print("请创建 .env 文件并设置这些变量")
    sys.exit(1)

# 设置模型变量
OPENROUTER_CLAUDE_MODEL = os.getenv("OPENROUTER_Claude_3.7_Sonnet")
OPENROUTER_GEMINI_MODEL = os.getenv("OPENROUTER_Gemini_2.5_Pro")
OPENROUTER_DS_R1_MODEL = os.getenv("OPENROUTER_DeepSeek_R1")

# 初始化API客户端
openrouter = OpenAI(
    api_key=os.getenv("OPENROUTER_API_KEY"),
    base_url=os.getenv("OPENROUTER_BASE_URL")
)
aihubmix = OpenAI(
    api_key=os.getenv("AIHUBMIX_API_KEY"),
    base_url=os.getenv("AIHUBMIX_BASE_URL")
)

# 初始化ARK客户端（可选）
ark_api_key = os.getenv("ARK_API_KEY")
ark_base_url = os.getenv("ARK_BASE_URL")
if ark_api_key and ark_base_url:
    ark_client = OpenAI(
        api_key=ark_api_key,
        base_url=ark_base_url
    )
    print("ARK图像生成服务已初始化")
else:
    ark_client = None
    print("ARK环境变量未配置，仅使用AIHUBMIX服务")

# 图像生成模型定义
IMAGE_MODELS = {
    "GPT-Image": "aihubmix",
    "Seedream": "ark"
}

def read_input_file(file_path):
    """读取输入文件内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        print(f"读取文件时出错: {e}")
        return None

def generate_content_with_title(content, output_dir=None, model=OPENROUTER_GEMINI_MODEL, reqs="", temperature=0.7):
    try:
        # 构建请求内容
        user_content = f"{content}\n\n额外要求：{reqs}" if reqs else content
            
        print("\n开始API调用...")
        response = openrouter.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_content}
            ],
            max_tokens=6000,
            temperature=temperature,
            response_format={"type": "json_object"}
        )
        
        response_content = response.choices[0].message.content.strip()
        print(f"API返回内容长度: {len(response_content)} 字符")
        
        # 解析JSON响应
        try:
            result = json.loads(response_content)
        except json.JSONDecodeError:
            print("JSON格式问题，尝试修复...")
            result = json.loads(json_repair.repair_json(response_content))
        
        print("成功解析JSON响应")
        
        # 保存JSON文件
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            json_path = os.path.join(output_dir, "content.json")
        else:
            json_path = "content.json"
            
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        
        print(f"JSON内容已保存到: {json_path}")
        return result
            
    except Exception as e:
        print(f"生成内容错误: {e}")
        return {"error": str(e)}

def generate_and_save_image(paragraph, output_dir, index, image_model="GPT-Image"):
    """基于段落生成图片并保存到指定目录"""
    
    if not paragraph:
        print("警告: 段落内容为空，无法生成图片")
        return None
        
    print(f"正在使用 {image_model} 生成图片 {index}...")
    
    # 确定使用的服务
    service = IMAGE_MODELS.get(image_model, "aihubmix")
    
    try:
        # API调用
        if service == "aihubmix":
            response = aihubmix.images.generate(
                model=os.getenv("AIHUBMIX_IMAGE_GENERATION_MODEL"),
                prompt=f"{paragraph}{IMAGE_STYLE_PROMPT}",
                n=1,
                size=IMAGE_SIZE,
                quality=IMAGE_QUALITY,
                moderation=IMAGE_MODERATION,
                background=IMAGE_BACKGROUND
            )
        elif service == "ark":
            if not ark_client:
                print("错误: ARK客户端未初始化，请检查环境变量配置")
                return None
            response = ark_client.images.generate(
                model=os.getenv("ARK_SeeDream_MODEL"),
                prompt=f"{paragraph}{IMAGE_STYLE_PROMPT}",
                size=IMAGE_SIZE,
                response_format="url"
            )
        else:
            print(f"错误: 未知的图像生成服务: {service}")
            return None
        
        if not response.data:
            print("错误: 未收到有效的图片数据")
            return None
        
        # 准备保存路径
        os.makedirs(output_dir, exist_ok=True)
        file_name = f"image_{index}.png"
        file_path = os.path.join(output_dir, file_name)
        if os.path.exists(file_path):
            file_name = f"image_{index}_{int(time.time())}.png"
            file_path = os.path.join(output_dir, file_name)
        
        # 保存图片
        if service == "aihubmix":
            image_data = response.data[0]
            image_bytes = base64.b64decode(image_data.b64_json)
            with open(file_path, "wb") as f:
                f.write(image_bytes)
        else:  # service == "ark"
            image_url = response.data[0].url
            img_response = requests.get(image_url)
            if img_response.status_code == 200:
                with open(file_path, "wb") as f:
                    f.write(img_response.content)
            else:
                print(f"图片下载失败: HTTP {img_response.status_code}")
                return None
        
        print(f"图片已保存: {file_name} (使用 {image_model})")
        return file_name
        
    except Exception as e:
        print(f"图片生成失败: {e}")
        return None

def create_docx_report_native(content_json, images, output_dir, image_mode=1, model_name=""):
    """创建Word文档报告"""
    doc = Document()
    
    # 设置文档默认字体，支持emoji显示
    doc.styles['Normal'].font.name = 'Apple Color Emoji'  # 支持emoji
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 改进的样式设置函数，支持emoji
    def add_styled_text(container, text, size=12, bold=False, font_name='微软雅黑'):
        run = container.add_run(text)
        # 设置支持emoji的字体
        run.font.name = font_name
        run.font.size = Pt(size)
        run.bold = bold
        # 设置字体支持emoji
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        # 添加emoji支持
        if any(ord(char) > 127 for char in text):  # 包含非ASCII字符（可能是emoji）
            run.font.name = 'Segoe UI Emoji'  # Windows emoji字体
            run._element.rPr.rFonts.set(qn('w:cs'), 'Segoe UI Emoji')
        return run
    
    # 添加文档标题
    heading = doc.add_heading(level=0)
    heading.clear()  # 清除默认内容
    add_styled_text(heading, "文章标题方案生成报告", size=20, bold=True)
    doc.add_paragraph()
    
    # 处理整篇文章配图（模式2）- 放在标题表格上面
    if image_mode == 2 and images and images[0]:
        image_path = os.path.join(output_dir, images[0])
        if os.path.exists(image_path):
            pic_paragraph = doc.add_paragraph()
            pic_paragraph.alignment = 1  # 居中
            pic_paragraph.add_run().add_picture(image_path, width=Inches(5.0))
            doc.add_paragraph()
    
    # 创建标题比较表格
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    # 表格标题
    headers = ["标题编号", "标题内容"]
    for i, header in enumerate(headers):
        cell_para = table.rows[0].cells[i].paragraphs[0]
        cell_para.clear()
        add_styled_text(cell_para, header, size=12, bold=True)
    
    # 添加三个标题方案
    title_names = ["标题一", "标题二", "标题三"]
    for i in range(1, 4):
        row = table.rows[i]
        # 标题编号
        cell_para_1 = row.cells[0].paragraphs[0]
        cell_para_1.clear()
        add_styled_text(cell_para_1, title_names[i-1], size=12)
        
        # 标题内容
        cell_para_2 = row.cells[1].paragraphs[0]
        cell_para_2.clear()
        add_styled_text(cell_para_2, content_json.get(f"title_{i}", "自媒体文章"), size=12)
    
    # 添加文章内容标题
    doc.add_paragraph()
    content_heading = doc.add_heading(level=0)
    content_heading.clear()
    add_styled_text(content_heading, "文章内容", size=18, bold=True)
    doc.add_paragraph()
    
    # 添加引言部分
    introduction = content_json.get("introduction", "")
    if introduction:
        intro_heading = doc.add_heading(level=1)
        intro_heading.clear()
        intro_heading.alignment = 1  # 居中对齐
        add_styled_text(intro_heading, "—— 引言 ——", size=14, bold=True)
        
        intro_para = doc.add_paragraph()
        add_styled_text(intro_para, introduction, size=12)
        doc.add_paragraph()
    
    # 添加段落内容
    paragraphs = content_json.get("paragraphs", [])
    for i, p in enumerate(paragraphs):
        subtitle = p.get("subtitle", "")
        content = p.get("content", "")
        
        # 添加子标题
        if subtitle:
            subheading = doc.add_heading(level=1)
            subheading.clear()
            subheading.alignment = 1  # 居中对齐
            add_styled_text(subheading, f"—— {subtitle} ——", size=14, bold=True)
        
        # 添加段落内容
        para = doc.add_paragraph()
        add_styled_text(para, content or "[内容缺失]", size=12)
        
        # 添加段落配图（模式1）
        if image_mode == 1 and i < len(images) and images[i]:
            image_path = os.path.join(output_dir, images[i])
            if os.path.exists(image_path):
                pic_paragraph = doc.add_paragraph()
                pic_paragraph.alignment = 1  # 居中
                pic_paragraph.add_run().add_picture(image_path, width=Inches(5.0))
                doc.add_paragraph()
    
    # 添加模型信息
    if model_name:
        doc.add_paragraph()
        model_para = doc.add_paragraph()
        model_para.alignment = 2  # 右对齐
        add_styled_text(model_para, f"Generated by: {model_name}", size=9)
    
    # 保存文档
    doc_path = os.path.join(output_dir, "report.docx")
    doc.save(doc_path)
    print(f"DOCX报告已生成: {doc_path}")

def validate_image_model(image_model):
    """
    验证图像模型的可用性。
    
    Args:
        image_model (str): 图像模型名称
        
    Returns:
        tuple: (is_valid, error_message)
    """
    if image_model not in IMAGE_MODELS:
        return False, f"不支持的图像模型 '{image_model}'，支持的模型: {', '.join(IMAGE_MODELS.keys())}"
    elif image_model == "Seedream" and not ark_client:
        return False, "Seedream模型需要配置ARK环境变量"
    return True, ""

def main(image_mode=1, model=OPENROUTER_GEMINI_MODEL, reqs="", temperature=0.7, image_model="Seedream"):
    # 读取输入文件
    content = read_input_file("input.txt")
    if not content:
        print("无法读取输入文件 input.txt")
        return
    
    # 验证图像模型选择和可用性
    is_valid, error_message = validate_image_model(image_model)
    if not is_valid:
        print(error_message)
        return
    
    # 创建输出目录
    timestamp = datetime.now().strftime("%m%d_%H%M")
    output_dir = f"content_{timestamp}"
    
    # 生成内容
    print(f"正在生成文章内容，使用模型：{model}")
    result = generate_content_with_title(content, output_dir, model=model, reqs=reqs, temperature=temperature)
    if "error" in result:
        print(f"内容生成失败: {result['error']}")
        return
    
    # 重命名目录为更友好的名称
    title = result.get("title_1", "未命名文章")
    clean_title = re.sub(r'[^\w\u4e00-\u9fff]', '', title)[:20]  # 保留中文和字母数字
    title_dir = f"{clean_title}_{timestamp}"
    if not os.path.exists(title_dir):
        os.rename(output_dir, title_dir)
        output_dir = title_dir
    
    # 生成图片
    paragraphs = result["paragraphs"]
    images = [None] * len(paragraphs)
    image_count = 0
    
    if image_mode == 1:  # 每段一张图
        print(f"正在使用 {image_model} 生成段落配图...")
        for i, p in enumerate(paragraphs):
            if p.get("content"):
                image_filename = generate_and_save_image(p["content"], output_dir, i + 1, image_model)
                if image_filename:
                    images[i] = image_filename
                    image_count += 1
    
    elif image_mode == 2:  # 整篇一张图
        print(f"正在使用 {image_model} 生成文章封面配图...")
        # 使用introduction和标题作为封面图生成依据
        title = result.get("title_1", "")
        introduction = result.get("introduction", "")
        cover_content = f"{title}。{introduction}"
        image_filename = generate_and_save_image(cover_content, output_dir, 0, image_model)
        if image_filename:
            images[0] = image_filename
            image_count = 1
    
    # 生成Word文档
    create_docx_report_native(result, images, output_dir, image_mode, model)
    
    # 输出结果
    print(f"\n生成完成! 保存至: {output_dir}")
    for i in range(1, 4):
        print(f"  标题{i}: {result.get(f'title_{i}', '无标题')}")
    print(f"  文件: report.docx, content.json, {image_count}张配图 (使用 {image_model})")

if __name__ == "__main__":
    """
    主程序入口
    使用示例：
        1. 为每个章节生成配图: main(image_mode=1)
        2. 为整篇文章生成一张配图: main(image_mode=2)
        3. 不生成任何配图: main(image_mode=3)
        
        使用不同模型:
        - 使用Claude 3.7 Sonnet: model=OPENROUTER_CLAUDE_MODEL
        - 使用DeepSeek R1: model=OPENROUTER_DS_R1_MODEL
        - 使用Gemini 2.5 Pro: model=OPENROUTER_GEMINI_MODEL
        
        使用不同图像生成模型:
        - 使用GPT-Image (GPT-4o): image_model="GPT-Image"
        - 使用Seedream (豆包的即梦3.0): image_model="Seedream"
        
        使用额外要求:
        - 在reqs参数中传入字符串，例如：
          main(reqs="文章风格要正式，使用专业术语")
        
        控制生成随机性:
        - temperature: 控制生成文本的随机性，范围0-1，值越大随机性越强，默认0.7
          例如: main(temperature=0.5)  # 更确定性的输出
               main(temperature=0.9)  # 更有创意的输出
    """
    
    main(
            image_mode=2, 
            model=OPENROUTER_GEMINI_MODEL, 
            reqs="文章风格生动活泼，像好友聊天", 
            temperature=0.7,
            image_model="Seedream"
        )