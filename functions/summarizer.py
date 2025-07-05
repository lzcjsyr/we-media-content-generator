#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
期刊摘要生成器

功能：
- 自动同步 awesome-english-ebooks 期刊仓库
- 从EPUB文件提取文本内容
- 使用AI生成期刊文章摘要
- 创建JSON摘要文件和Word报告

支持的期刊：
- 经济学人 (The Economist)
- 纽约客 (The New Yorker)  
- 大西洋月刊 (The Atlantic)
- 连线 (Wired)

使用方法：
    python summarizer.py [magazine_type] [-o output_dir] [--skip-git-check]
    
    示例：
    python summarizer.py economist          # 处理经济学人
    python summarizer.py 1                  # 处理第1个杂志（经济学人）
    python summarizer.py                    # 交互式选择

输出：
    在 摘要汇总/ 目录下生成：
    - JSON摘要文件
    - Word汇总报告
"""

import os
import glob
import json
import re
import sys
import argparse
import subprocess
import docx
from docx.oxml.ns import qn
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
from openai import OpenAI
from dotenv import load_dotenv

# 加载环境变量
load_dotenv()

# Git仓库路径 - 使用相对路径
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
# 从 functions/ 向上一级到 内容生成器/，再向上一级到 期刊系列内容生成/，然后进入 awesome-english-ebooks/
REPO_PATH = os.path.join(os.path.dirname(os.path.dirname(SCRIPT_DIR)), "awesome-english-ebooks")

# 杂志配置信息
MAGAZINE_CONFIG = {
    "economist": {"name": "The Economist", "base_dir": "01_economist", "folder_pattern": "te_*", "file_pattern": "*.epub", "title": "经济学人"},
    "new_yorker": {"name": "The New Yorker", "base_dir": "02_new_yorker", "folder_pattern": "20*", "file_pattern": "*.epub", "title": "纽约客"},
    "atlantic": {"name": "The Atlantic", "base_dir": "04_atlantic", "folder_pattern": "20*", "file_pattern": "*.epub", "title": "大西洋月刊"},
    "wired": {"name": "Wired", "base_dir": "05_wired", "folder_pattern": "20*", "file_pattern": "*.epub", "title": "连线"}
}

# 使用OpenRouter API初始化OpenAI客户端
client = OpenAI(
    base_url=os.getenv("OPENROUTER_BASE_URL"),
    api_key=os.getenv("OPENROUTER_API_KEY")
)

def check_and_update_repo():
    """检查并更新或克隆 awesome-english-ebooks 期刊仓库"""
    repo_url = "https://github.com/hehonghui/awesome-english-ebooks.git"
    
    def run_git_cmd(cmd, timeout=30, cwd=None):
        """运行Git命令的辅助函数"""
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout, cwd=cwd or REPO_PATH)
            return result.returncode == 0, result.stdout.strip(), result.stderr
        except subprocess.TimeoutExpired:
            return False, "", "操作超时"
        except Exception as e:
            return False, "", str(e)
    
    print(f"正在检查期刊仓库状态 ({REPO_PATH})...")
    
    # 检查仓库是否存在
    if not os.path.exists(REPO_PATH):
        print("未发现期刊仓库，正在克隆...")
        parent_dir = os.path.dirname(REPO_PATH)
        os.makedirs(parent_dir, exist_ok=True)
        
        success, output, error = run_git_cmd(['git', 'clone', repo_url, "awesome-english-ebooks"], 120, parent_dir)
        if success:
            print("✓ 期刊仓库克隆成功!")
            return True
        else:
            print(f"✗ 期刊仓库克隆失败: {error}")
            return False
    
    # 检查是否为有效Git仓库
    if not os.path.exists(os.path.join(REPO_PATH, '.git')):
        print(f"⚠ 发现 {REPO_PATH} 目录，但不是有效的Git仓库")
        return False
    
    # 获取远程更新信息
    success, _, error = run_git_cmd(['git', 'fetch'])
    if not success:
        print(f"获取远程信息失败: {error}")
        return False
    
    # 获取当前分支和提交信息
    success, current_branch, _ = run_git_cmd(['git', 'rev-parse', '--abbrev-ref', 'HEAD'])
    if not success:
        print("获取分支信息失败")
        return False
    
    # 比较本地和远程提交
    success, local_commit, _ = run_git_cmd(['git', 'rev-parse', 'HEAD'])
    success2, remote_commit, _ = run_git_cmd(['git', 'rev-parse', f'origin/{current_branch}'])
    
    if not (success and success2):
        print("获取提交信息失败")
        return False
    
    if local_commit == remote_commit:
        print("✓ 期刊仓库已经是最新版本")
        return True
    
    # 检查并更新
    success, status_output, _ = run_git_cmd(['git', 'status', '-uno'])
    if not success:
        print("检查状态失败")
        return False
    
    if "Your branch is behind" in status_output:
        print("发现期刊仓库有更新，正在执行 git pull...")
        success, output, error = run_git_cmd(['git', 'pull', '--stat'], 60)
        if success:
            print("✓ 期刊仓库更新成功!")
            
            # 显示更新的期刊文件夹
            if output.strip():
                updated_folders = set()
                for line in output.strip().split('\n'):
                    if '|' in line and ('+' in line or '-' in line):
                        file_path = line.split('|')[0].strip()
                        if file_path and not file_path.startswith(('create mode', 'delete mode')):
                            path_parts = file_path.split('/')
                            if len(path_parts) >= 2:
                                folder_path = '/'.join(path_parts[:2])
                                updated_folders.add(folder_path)
                
                if updated_folders:
                    print("📂 更新的期刊文件夹:")
                    for folder in sorted(updated_folders):
                        print(f"   📁 {folder}")
            
            return True
        else:
            print(f"✗ 期刊仓库更新失败: {error}")
            return False
    elif "have diverged" in status_output:
        print("检测到分支分叉，正在强制更新到远程版本...")
        success, _, error = run_git_cmd(['git', 'reset', '--hard', f'origin/{current_branch}'])
        if success:
            print("✓ 期刊仓库已强制更新到远程版本!")
            return True
        else:
            print(f"✗ 强制更新失败: {error}")
            return False
    else:
        print("✓ 检测到更新但状态复杂，建议手动处理")
        return False

def extract_text_from_epub(epub_path):
    """使用ebooklib和BeautifulSoup从EPUB文件中提取文本。"""
    try:
        text_parts = []
        book = epub.read_epub(epub_path)
        
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                try:
                    content = item.get_content().decode('utf-8')
                    soup = BeautifulSoup(content, 'html.parser')
                    
                    # 移除script和style标签
                    for script in soup(["script", "style"]):
                        script.decompose()
                    
                    # 提取文本
                    text = soup.get_text(' ', strip=True)
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    print(f"  处理EPUB章节时出错: {e}")
                    continue
        
        return '\n\n'.join(text_parts)
    except Exception as e:
        print(f"从{epub_path}提取文本时出错: {e}")
        return ""

def parse_date_from_filename(filename, magazine_type):
    """从文件名中解析日期"""
    match = re.search(r'(\d{4})\.(\d{2})\.(\d{2})', filename)
    return f"{match.group(1)}-{match.group(2)}-{match.group(3)}" if match else None

def summarize_with_llm(epub_text, journal_name, publication_date):
    """使用LLM生成杂志摘要"""
    prompt = f"""请分析以下《{journal_name}》杂志（发行日期：{publication_date}）的内容。
请提取出杂志中每篇文章的主要内容，每篇文章的总结约300字，用简体中文输出。

请严格按照以下JSON结构输出：
{{
  "journal_name": "{journal_name}",
  "publication_date": "{publication_date}",
  "articles": [
    {{
      "title": "标题示例1",
      "Chinese_title": "对应的中文标题示例1",
      "summary": "文章摘要内容..."
    }}
  ]
}}

重要说明：
1. "publication_date" 字段必须严格使用 YYYY-MM-DD 格式（如 2024-01-02）
2. 不要使用月份名称（如 January/February 2024）或其他日期格式
3. 请使用传入的 publication_date 参数值：{publication_date}

以下是杂志内容：
{epub_text}

请只输出有效的JSON格式。"""

    try:
        response = client.chat.completions.create(
            model=os.getenv("OPENROUTER_Gemini_2.5_Pro"),
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            max_tokens=30000,
            temperature=0.3
        )
        
        content = response.choices[0].message.content.strip()
        if not content:
            raise ValueError("LLM返回了空响应")
            
        # 清理可能的markdown代码块标记
        if content.startswith('```json'):
            content = content[7:]
        if content.endswith('```'):
            content = content[:-3]
        content = content.strip()
        
        data = json.loads(content)
        
        # 验证响应格式
        required_keys = ('journal_name', 'publication_date', 'articles')
        if not (isinstance(data, dict) and all(k in data for k in required_keys) and
               isinstance(data.get('articles'), list) and len(data['articles']) > 0):
            raise ValueError("LLM响应格式无效")
            
        return data
        
    except (json.JSONDecodeError, ValueError) as e:
        print(f"处理LLM响应时出错: {e}")
        print(f"响应内容前500字符: {content[:500] if 'content' in locals() else 'None'}")
    except Exception as e:
        print(f"调用LLM时出错: {e}")
    return None

def save_json_summary(summary_data, output_dir, filename):
    """将摘要保存为JSON文件。"""
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(summary_data, f, ensure_ascii=False, indent=2)
    
    return output_path

def create_docx_report(json_files, output_path, magazine_title="杂志"):
    """创建Word报告"""
    # 按出版日期排序文件
    sorted_files = sorted(json_files, key=lambda x: json.load(open(x, 'r', encoding='utf-8')).get('publication_date', ''))
    
    doc = docx.Document()
    
    # 设置字体样式
    def set_font(style_name, font_name='SimSun', font_size=10.5):
        style = doc.styles[style_name]
        style.font.name = font_name
        style.font.size = docx.shared.Pt(font_size)
        style.font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    for style, size in [('Normal', 10.5), ('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)]:
        set_font(style, 'SimSun', size)
    
    doc.add_heading(f'《{magazine_title}》期刊摘要汇总', 0)
    
    # 处理每个JSON文件
    for json_file in sorted_files:
        try:
            with open(json_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            doc.add_heading(f"{data.get('journal_name', '未知')} ({data.get('publication_date', '未知')})", 1)
            
            for article in data.get('articles', []):
                english_title = article.get('title', '未知标题')
                chinese_title = article.get('Chinese_title', '')
                full_title = f"{chinese_title} ({english_title})" if chinese_title else english_title
                
                doc.add_heading(full_title, 2)
                doc.add_paragraph(article.get('summary', '无可用摘要'))
            
            doc.add_page_break()
            
        except Exception as e:
            print(f"处理{json_file}时出错: {e}")
    
    doc.save(output_path)
    print(f"Word报告已保存至{output_path}")

def _get_magazine_paths(magazine_type, base_output_dir):
    """获取杂志相关的路径和文件信息"""
    config = MAGAZINE_CONFIG[magazine_type]
    base_path = os.path.join(REPO_PATH, config['base_dir'])
    
    # 修复输出目录路径计算
    if os.path.isabs(base_output_dir):
        magazine_output_dir = os.path.join(base_output_dir, config['title'])
    else:
        # 修复路径：从functions目录向上一级到项目根目录
        project_root = os.path.dirname(SCRIPT_DIR)  # 向上一级到项目根目录
        magazine_output_dir = os.path.join(project_root, base_output_dir, config['title'])
    
    # 确保输出目录存在
    os.makedirs(magazine_output_dir, exist_ok=True)
    
    # 查找所有符合模式的目录
    magazine_dirs = glob.glob(os.path.join(base_path, config['folder_pattern']))
    
    # 收集所有EPUB文件
    all_epub_files = []
    valid_epub_files = []
    valid_folders = []
    
    # 统计有效的文件夹和EPUB文件
    for dir_path in magazine_dirs:
        if os.path.isdir(dir_path):
            epub_files = glob.glob(os.path.join(dir_path, config['file_pattern']))
            if epub_files:  # 只统计包含EPUB文件的文件夹
                valid_folders.append(dir_path)
                all_epub_files.extend(epub_files)
    
    # 处理EPUB文件，提取日期和生成预期的JSON路径，去重处理
    epub_info = []
    unique_dates = set()
    seen_dates = {}  # 用于跟踪已见过的日期
    
    for epub_path in all_epub_files:
        filename = os.path.basename(epub_path)
        publication_date = parse_date_from_filename(filename, magazine_type)
        if publication_date:
            unique_dates.add(publication_date)
            # 对重复日期进行去重，保留第一个找到的文件
            if publication_date not in seen_dates:
                json_filename = f"{config['name']}_{publication_date.replace('-', '')}.json"
                json_path = os.path.join(magazine_output_dir, json_filename)
                epub_info.append({
                    'epub_path': epub_path,
                    'json_path': json_path,
                    'publication_date': publication_date,
                    'filename': filename
                })
                valid_epub_files.append(epub_path)
                seen_dates[publication_date] = True
    
    # 获取已存在的JSON文件
    existing_json_files = []
    if os.path.exists(magazine_output_dir):
        existing_json_files = [f for f in glob.glob(os.path.join(magazine_output_dir, "*.json")) 
                              if not f.endswith('摘要汇总.docx')]  # 排除Word文档
    existing_json_set = set(existing_json_files)
    
    return {
        'config': config,
        'magazine_dirs': magazine_dirs,
        'all_epub_files': all_epub_files,
        'valid_epub_files': valid_epub_files,
        'valid_folders': valid_folders,
        'epub_info': epub_info,
        'unique_dates': unique_dates,
        'existing_json_files': existing_json_files,
        'existing_json_set': existing_json_set,
        'magazine_output_dir': magazine_output_dir
    }

def count_magazine_files(magazine_type, base_output_dir="摘要汇总"):
    """统计杂志文件数量"""
    try:
        paths = _get_magazine_paths(magazine_type, base_output_dir)
        existing_expected = {info['json_path'] for info in paths['epub_info']} & set(paths['existing_json_files'])
        
        return {
            "total_folders": len(paths['valid_folders']),
            "unique_dates": len(paths['unique_dates']),
            "existing_json": len(paths['existing_json_files']),
            "to_process": len(paths['valid_folders']) - len(existing_expected)
        }
    except Exception as e:
        print(f"统计杂志文件时出错: {e}")
        return {"total_folders": 0, "unique_dates": 0, "existing_json": 0, "to_process": 0}

def process_magazine(magazine_type, base_output_dir="摘要汇总"):
    """处理指定类型的杂志"""
    if magazine_type not in MAGAZINE_CONFIG:
        print(f"错误：未知的杂志类型 '{magazine_type}'")
        return
    
    config = MAGAZINE_CONFIG[magazine_type]
    print(f"\n{'-' * 20} 开始处理 {config['title']} {'-' * 20}")
    
    try:
        paths = _get_magazine_paths(magazine_type, base_output_dir)
        existing_expected = {info['json_path'] for info in paths['epub_info']} & paths['existing_json_set']
        needs_processing = [info for info in paths['epub_info'] if info['json_path'] not in paths['existing_json_set']]
        
        # 打印统计信息
        print(f"\n{'=' * 60}")
        print(f"【{config['name']}摘要统计信息】")
        print(f"  - 总期数: {len(paths['valid_folders'])} 期")
        print(f"  - 已完成摘要: {len(existing_expected)} 期")
        print(f"  - 待处理: {len(needs_processing)} 期")
        print("=" * 60 + "\n")
        
        if not needs_processing:
            print("所有文件已经处理完毕，无需进行新的摘要处理\n")
            # 如果没有需要处理的文件，仍然生成Word报告（如果有现有的JSON文件）
            word_report_success = False
            if existing_expected:
                try:
                    docx_filename = f"《{config['title']}》摘要汇总.docx"
                    docx_path = os.path.join(paths['magazine_output_dir'], docx_filename)
                    create_docx_report(list(existing_expected), docx_path, config['title'])
                    print(f"Word报告已保存至{docx_path}")
                    word_report_success = True
                except Exception as e:
                    print(f"创建Word文档时出错: {e}")
            
            return {
                'total': 0,
                'successful': 0,
                'failed': 0,
                'word_report': word_report_success,
                'all_successful': True,
                'partially_successful': True,
                'already_completed': True
            }
        else:
            print("开始进行摘要处理...\n")
        
        # 处理需要生成摘要的文件
        json_files = list(existing_expected)
        successful_count = 0
        failed_count = 0
        
        for i, info in enumerate(needs_processing):
            try:
                print(f"[{i+1}/{len(needs_processing)}] 正在处理{info['filename']}...")
                
                epub_text = extract_text_from_epub(info['epub_path'])
                if not epub_text:
                    print(f"  从{info['epub_path']}提取文本失败，跳过")
                    failed_count += 1
                    continue
                
                print("  正在使用Gemini 2.5 Pro进行摘要...")
                summary_data = summarize_with_llm(epub_text, config['name'], info['publication_date'])
                
                if summary_data:
                    save_json_summary(summary_data, paths['magazine_output_dir'], os.path.basename(info['json_path']))
                    print(f"  摘要已保存至{info['json_path']}")
                    json_files.append(info['json_path'])
                    successful_count += 1
                else:
                    print(f"  为{info['filename']}生成摘要失败，跳过")
                    failed_count += 1
            except Exception as e:
                print(f"  处理{info.get('filename', '未知文件')}时出错: {e}")
                failed_count += 1
                continue
        
        # 创建Word报告
        word_report_success = False
        if json_files:
            try:
                docx_filename = f"《{config['title']}》摘要汇总.docx"
                docx_path = os.path.join(paths['magazine_output_dir'], docx_filename)
                create_docx_report(json_files, docx_path, config['title'])
                print(f"Word报告已保存至{docx_path}")
                word_report_success = True
            except Exception as e:
                print(f"创建Word文档时出错: {e}")
        
        # 打印处理结果汇总
        total_processed = len(needs_processing)
        print(f"\n{'='*50}")
        print(f"📊 {config['title']} 处理结果汇总:")
        print(f"  📝 总计处理: {total_processed} 篇")
        print(f"  ✅ 成功生成: {successful_count} 篇")
        print(f"  ❌ 处理失败: {failed_count} 篇")
        if word_report_success:
            print(f"  📄 Word报告: 已生成")
        
        if failed_count > 0:
            print(f"\n⚠️  部分文件处理失败，这通常是由于LLM响应不稳定导致的。")
            print(f"   建议重新运行程序处理失败的文件。")
        
        print(f"{'='*50}")
        
        # 返回处理结果
        return {
            'total': total_processed,
            'successful': successful_count,
            'failed': failed_count,
            'word_report': word_report_success,
            'all_successful': failed_count == 0,
            'partially_successful': successful_count > 0
        }
        
    except Exception as e:
        print(f"处理杂志{magazine_type}时发生错误: {e}")
        return {
            'total': 0,
            'successful': 0,
            'failed': 0,
            'word_report': False,
            'all_successful': False,
            'partially_successful': False,
            'error': str(e)
        }

def main():
    """主函数，处理命令行参数"""
    parser = argparse.ArgumentParser(description="杂志摘要生成工具")
    
    # 添加杂志类型参数
    parser.add_argument(
        "magazine", 
        nargs="?",  # 可选参数
        choices=list(MAGAZINE_CONFIG.keys()) + [str(i) for i in range(1, len(MAGAZINE_CONFIG) + 1)],
        help="要处理的杂志类型或编号"
    )
    
    # 添加输出目录参数
    parser.add_argument(
        "-o", "--output",
        default="摘要汇总",
        help="摘要输出的基础目录 (默认: 摘要汇总)"
    )
    
    # 添加跳过Git检查的选项
    parser.add_argument(
        "--skip-git-check",
        action="store_true",
        help="跳过期刊仓库更新检查"
    )
    
    args = parser.parse_args()
    
    # 检查并更新期刊仓库（除非用户选择跳过）
    if not args.skip_git_check:
        print("=" * 60)
        print("【期刊仓库更新检查】")
        print("=" * 60)
        
        # 保存当前工作目录
        original_cwd = os.getcwd()
        
        try:
            if not check_and_update_repo():
                print("\n警告：期刊仓库更新检查失败，但程序将继续运行...")
                print("如果需要跳过检查，请使用 --skip-git-check 参数\n")
            else:
                print()
        finally:
            # 恢复原始工作目录
            os.chdir(original_cwd)
    else:
        print("已跳过期刊仓库更新检查\n")
    
    # 显示文件统计信息
    print("正在收集文件统计信息...") 
    
    stats = {}
    for magazine_type, config in MAGAZINE_CONFIG.items():
        try:
            stats[magazine_type] = count_magazine_files(magazine_type, args.output)
        except Exception as e:
            print(f"统计 {config['title']} 信息时出错: {e}")
            stats[magazine_type] = {"total_folders": 0, "unique_dates": 0, "existing_json": 0, "to_process": 0}
    
    # 计算中文字符显示宽度的函数
    def display_width(text):
        """计算字符串在终端中的显示宽度"""
        return sum(2 if '\u4e00' <= char <= '\u9fff' else 1 for char in text)
    
    def format_column(text, width):
        """格式化列，考虑中文字符宽度"""
        text_str = str(text)
        padding = width - display_width(text_str)
        return text_str + ' ' * max(0, padding)
    
    # 表格显示
    divider = '-' * 60
    print(f"\n{divider}")
    
    # 表头
    header = (format_column('期刊名称', 16) + 
             format_column('总期数', 10) + 
             format_column('已生成', 10) + 
             format_column('待处理', 10))
    print(header)
    print(divider)
    
    # 数据行
    for magazine_type, config in MAGAZINE_CONFIG.items():
        s = stats[magazine_type]
        row = (format_column(config['title'], 16) + 
               format_column(str(s['total_folders']), 10) + 
               format_column(str(s['existing_json']), 10) + 
               format_column(str(s['to_process']), 10))
        print(row)
    
    print(f"{divider}\n")
    
    # 处理杂志选择
    if args.magazine:
        if args.magazine.isdigit():
            choice = int(args.magazine)
            if 1 <= choice <= len(MAGAZINE_CONFIG):
                selected_magazine = list(MAGAZINE_CONFIG.keys())[choice-1]
            else:
                print(f"错误: 无效的杂志编号 {choice}")
                sys.exit(1)
        elif args.magazine in MAGAZINE_CONFIG:
            selected_magazine = args.magazine
        else:
            print(f"错误: 未知的杂志类型 '{args.magazine}'")
            sys.exit(1)
    else:
        # 交互式选择
        print("请选择要处理的杂志类型:")
        for i, (key, config) in enumerate(MAGAZINE_CONFIG.items(), 1):
            print(f"{i}. {config['name']} ({config['title']})")
        
        while True:
            try:
                choice = input(f"请输入选项编号 (1-{len(MAGAZINE_CONFIG)}), 或按 Ctrl+C 退出: ")
                if choice.isdigit() and 1 <= int(choice) <= len(MAGAZINE_CONFIG):
                    selected_magazine = list(MAGAZINE_CONFIG.keys())[int(choice)-1]
                    break
                print(f"错误: 请输入1到{len(MAGAZINE_CONFIG)}之间的数字")
            except KeyboardInterrupt:
                print("\n操作已取消")
                sys.exit(0)
    
    # 处理选定的杂志
    process_magazine(selected_magazine, args.output)

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n程序已被用户中断")
        sys.exit(0)
