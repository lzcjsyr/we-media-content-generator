#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import glob
import json
import re
import sys
import argparse
from datetime import datetime
import docx
from docx.oxml.ns import qn
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
from openai import OpenAI
from dotenv import load_dotenv

# 加载环境变量
load_dotenv()

# 杂志配置信息
MAGAZINE_CONFIG = {
    "economist": {
        "name": "The Economist",
        "base_dir": "01_economist",
        "folder_pattern": "te_*",
        "file_pattern": "*.epub",
        "title": "经济学人"
    },
    "new_yorker": {
        "name": "The New Yorker",
        "base_dir": "02_new_yorker",
        "folder_pattern": "20*",  # 年份开头的文件夹
        "file_pattern": "*.epub",
        "title": "纽约客"
    },
    "atlantic": {
        "name": "The Atlantic",
        "base_dir": "04_atlantic",
        "folder_pattern": "20*",  # 年份开头的文件夹
        "file_pattern": "*.epub",
        "title": "大西洋月刊"
    },
    "wired": {
        "name": "Wired",
        "base_dir": "05_wired",
        "folder_pattern": "20*",  # 年份开头的文件夹
        "file_pattern": "*.epub",
        "title": "连线"
    }
}

# 使用OpenRouter API初始化OpenAI客户端
client = OpenAI(
    base_url=os.getenv("OPENROUTER_BASE_URL"),
    api_key=os.getenv("OPENROUTER_API_KEY")
)

def extract_text_from_epub(epub_path):
    """使用ebooklib和BeautifulSoup从EPUB文件中提取文本。"""
    try:
        text_parts = []
        book = epub.read_epub(epub_path)
        
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                try:
                    content = item.get_content().decode('utf-8')
                    soup = BeautifulSoup(content, 'html.parser')
                    
                    # 移除script和style标签
                    for script in soup(["script", "style"]):
                        script.decompose()
                    
                    # 提取文本
                    text = soup.get_text(' ', strip=True)
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    print(f"  处理EPUB章节时出错: {e}")
                    continue
        
        return '\n\n'.join(text_parts)
    except Exception as e:
        print(f"从{epub_path}提取文本时出错: {e}")
        return ""

def parse_date_from_filename(filename, magazine_type):

    # 不同杂志可能有不同的文件名格式
    if magazine_type == "economist":
        match = re.search(r'(\d{4})\.(\d{2})\.(\d{2})', filename)
    elif magazine_type in ["new_yorker", "atlantic", "wired"]:
        match = re.search(r'(\d{4})\.(\d{2})\.(\d{2})', filename)
    else:
        match = re.search(r'(\d{4})\.(\d{2})\.(\d{2})', filename)
        
    if match:
        year, month, day = match.groups()
        return f"{year}-{month}-{day}"
    return None

def summarize_with_llm(epub_text, journal_name, publication_date):
    
    prompt = f"""
请分析以下《{journal_name}》杂志（发行日期：{publication_date}）的内容。
请提取出杂志中每篇文章的主要内容，每篇文章的总结约300字，用简体中文输出。

请严格按照以下JSON结构输出，不要增加或删除字段：

```json
{{
  "journal_name": "{journal_name}",
  "publication_date": "{publication_date}",
  "articles": [
    {{
      "title": "文章标题示例1",
      "summary": "这是第一篇文章的摘要示例，摘要应该简洁明确，避免使用复杂或专业的术语，除非这些术语是理解文章的关键。在撰写摘要时，应保持客观，准确反映原文的观点，而不是提供个人的观点或评价。最后，摘要应该结构良好，有逻辑性，以便读者能易于理解和跟进文章的内容和论点。"
    }}
  ]
}}
```

以下是杂志内容，请分析并生成摘要：
{epub_text}

请只输出有效的JSON格式，不要加入其他说明或注释。
"""

    try:
        response = client.chat.completions.create(
            model=os.getenv("OPENROUTER_Gemini_2.5_Pro"),
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            max_tokens=30000,
            temperature=0.3
        )
        
        if not (content := response.choices[0].message.content.strip()):
            raise ValueError("LLM返回了空响应")
            
        data = json.loads(content)
        
        if not (isinstance(data, dict) and 
               all(k in data for k in ('journal_name', 'publication_date', 'articles')) and
               isinstance(data.get('articles'), list) and 
               len(data['articles']) > 0):
            raise ValueError("LLM响应格式无效")
            
        return data
        
    except (json.JSONDecodeError, ValueError) as e:
        print(f"处理LLM响应时出错: {e}")
    except Exception as e:
        print(f"调用LLM时出错: {e}")
    return None

def save_json_summary(summary_data, output_dir, filename):
    """将摘要保存为JSON文件。"""
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(summary_data, f, ensure_ascii=False, indent=2)
    
    return output_path

def create_docx_report(json_files, output_path, magazine_title="杂志"):

    # 按出版日期排序文件
    sorted_files = sorted(json_files, key=lambda x: json.load(open(x, 'r', encoding='utf-8')).get('publication_date', ''))
    
    # 创建新文档
    doc = docx.Document()
    
    # 设置文档默认字体为宋体
    def set_font(style_name, font_name='SimSun', font_size=10.5):
        style = doc.styles[style_name]
        font = style.font
        font.name = font_name
        font.size = docx.shared.Pt(font_size)
        # 设置中文字体
        font._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    # 设置各种样式的字体
    set_font('Normal', 'SimSun', 10.5)  # 正文
    set_font('Heading 1', 'SimSun', 16)  # 一级标题
    set_font('Heading 2', 'SimSun', 14)  # 二级标题
    set_font('Heading 3', 'SimSun', 12)  # 三级标题
    
    # 添加主标题
    doc.add_heading(f'《{magazine_title}》期刊摘要汇总', 0).style = 'Heading 1'
    
    # 从每个JSON文件添加内容
    for json_file in sorted_files:
        try:
            with open(json_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # 添加期刊信息
            doc.add_heading(f"{data.get('journal_name', '未知')} ({data.get('publication_date', '未知')})", 1)
            
            # 添加文章
            for article in data.get('articles', []):
                doc.add_heading(article.get('title', '未知标题'), 2)
                doc.add_paragraph(article.get('summary', '无可用摘要'))
            
            # 在期刊之间添加分页符
            doc.add_page_break()
            
        except Exception as e:
            print(f"处理{json_file}时出错: {e}")
    
    # 保存文档
    doc.save(output_path)
    print(f"Docx报告已保存至{output_path}")

def _get_magazine_paths(magazine_type, base_output_dir):
    """获取杂志相关的路径和文件信息"""
    config = MAGAZINE_CONFIG[magazine_type]
    base_path = f"/Users/dielangli/Documents/文件/Books/awesome-english-ebooks/{config['base_dir']}"
    magazine_output_dir = os.path.join(base_output_dir, config['title'])
    
    # 确保输出目录存在
    os.makedirs(magazine_output_dir, exist_ok=True)
    
    # 查找所有符合模式的目录
    magazine_dirs = glob.glob(os.path.join(base_path, config['folder_pattern']))
    
    # 收集所有EPUB文件
    all_epub_files = []
    valid_epub_files = []
    valid_folders = []
    
    # 统计有效的文件夹和EPUB文件
    for dir_path in magazine_dirs:
        if os.path.isdir(dir_path):
            epub_files = glob.glob(os.path.join(dir_path, config['file_pattern']))
            if epub_files:  # 只统计包含EPUB文件的文件夹
                valid_folders.append(dir_path)
                all_epub_files.extend(epub_files)
    
    # 处理EPUB文件，提取日期和生成预期的JSON路径
    epub_info = []
    unique_dates = set()
    for epub_path in all_epub_files:
        filename = os.path.basename(epub_path)
        publication_date = parse_date_from_filename(filename, magazine_type)
        if publication_date:
            unique_dates.add(publication_date)
            json_filename = f"{config['name']}_{publication_date.replace('-', '')}.json"
            json_path = os.path.join(magazine_output_dir, json_filename)
            epub_info.append({
                'epub_path': epub_path,
                'json_path': json_path,
                'publication_date': publication_date,
                'filename': filename
            })
            valid_epub_files.append(epub_path)
    
    # 获取已存在的JSON文件
    existing_json_files = glob.glob(os.path.join(magazine_output_dir, "*.json")) if os.path.exists(magazine_output_dir) else []
    existing_json_set = set(existing_json_files)
    
    return {
        'config': config,
        'magazine_dirs': magazine_dirs,
        'all_epub_files': all_epub_files,  # 所有找到的EPUB文件
        'valid_epub_files': valid_epub_files,  # 有效的EPUB文件（能解析出日期的）
        'valid_folders': valid_folders,  # 包含EPUB文件的文件夹
        'epub_info': epub_info,
        'unique_dates': unique_dates,
        'existing_json_files': existing_json_files,
        'existing_json_set': existing_json_set,
        'magazine_output_dir': magazine_output_dir
    }

def count_magazine_files(magazine_type, base_output_dir="summary"):
    try:
        # 获取杂志的所有路径信息
        paths = _get_magazine_paths(magazine_type, base_output_dir)
        
        # 计算已存在的JSON文件中，哪些是预期应该存在的（即有对应的EPUB文件）
        # 1. 从epub_info中提取所有预期的JSON路径
        # 2. 与已存在的JSON文件路径取交集
        existing_expected = {
            info['json_path'] for info in paths['epub_info']
        } & set(paths['existing_json_files'])
        
        return {
            # 包含EPUB文件的文件夹数量（即总期数）
            "total_folders": len(paths['valid_folders']),
            # 唯一日期数量（用于去重）
            "unique_dates": len(paths['unique_dates']),
            # 已存在的JSON文件数量
            "existing_json": len(paths['existing_json_files']),
            # 需要处理的文件数 = 总期数(有效文件夹数) - 已存在的有效JSON文件数
            "to_process": len(paths['valid_folders']) - len(existing_expected)
        }
    except Exception as e:
        print(f"统计杂志文件时出错: {e}")
        return {
            "total_folders": 0,
            "unique_dates": 0, 
            "existing_json": 0, 
            "to_process": 0
        }

def process_magazine(magazine_type, base_output_dir="summary"):
    """处理指定类型的杂志"""
    if magazine_type not in MAGAZINE_CONFIG:
        print(f"错误：未知的杂志类型 '{magazine_type}'")
        return
    
    config = MAGAZINE_CONFIG[magazine_type]
    print(f"\n{'-' * 20} 开始处理 {config['title']} {'-' * 20}")
    
    try:
        # 获取杂志相关路径和文件信息
        paths = _get_magazine_paths(magazine_type, base_output_dir)
        
        # 计算需要处理的文件
        existing_expected = {info['json_path'] for info in paths['epub_info']} & paths['existing_json_set']
        needs_processing = [info for info in paths['epub_info'] if info['json_path'] not in paths['existing_json_set']]
        
        # 打印统计信息
        print(f"\n{'=' * 60}")
        print(f"【{config['name']}摘要统计信息】")
        print(f"  - 总期数: {len(paths['valid_folders'])} 期")
        print(f"  - 已完成摘要: {len(existing_expected)} 期")
        print(f"  - 待处理: {len(needs_processing)} 期")
        
        if not needs_processing:
            print("\n所有文件已经处理完毕，无需进行新的摘要处理\n")
        else:
            print(f"\n开始进行摘要处理...\n")
        print("=" * 60 + "\n")
        
        # 处理需要生成摘要的文件
        json_files = list(existing_expected)
        for i, info in enumerate(needs_processing):
            try:
                print(f"[{i+1}/{len(needs_processing)}] 正在处理{info['filename']}...")
                
                # 从EPUB提取文本
                epub_text = extract_text_from_epub(info['epub_path'])
                if not epub_text:
                    print(f"  从{info['epub_path']}提取文本失败，跳过")
                    continue
                
                # 使用LLM进行摘要
                print(f"  正在使用Gemini 2.5 Pro进行摘要...")
                summary_data = summarize_with_llm(
                    epub_text, 
                    config['name'], 
                    info['publication_date']
                )
                
                # 保存摘要结果
                if summary_data is not None:
                    save_json_summary(summary_data, paths['magazine_output_dir'], os.path.basename(info['json_path']))
                    print(f"  摘要已保存至{info['json_path']}")
                    json_files.append(info['json_path'])
                else:
                    print(f"  为{info['filename']}生成摘要失败，跳过")
            except Exception as e:
                print(f"  处理{info.get('filename', '未知文件')}时出错: {str(e)}")
                continue
        
        # 从所有JSON文件创建docx报告
        if json_files:
            try:
                docx_filename = f"《{config['title']}》摘要汇总.docx"
                docx_path = os.path.join(paths['magazine_output_dir'], docx_filename)
                create_docx_report(json_files, docx_path, config['title'])
            except Exception as e:
                print(f"创建Word文档时出错: {str(e)}")
    except Exception as e:
        print(f"处理杂志{magazine_type}时发生错误: {str(e)}")

def main():
    """主函数，处理命令行参数"""
    parser = argparse.ArgumentParser(description="杂志摘要生成工具")
    
    # 添加杂志类型参数
    parser.add_argument(
        "magazine", 
        nargs="?",  # 可选参数
        choices=list(MAGAZINE_CONFIG.keys()) + [str(i) for i in range(1, len(MAGAZINE_CONFIG) + 1)],
        help="要处理的杂志类型或编号"
    )
    
    # 添加输出目录参数
    parser.add_argument(
        "-o", "--output",
        default="summary",
        help="摘要输出的基础目录 (默认: summary)"
    )
    
    args = parser.parse_args()
    
    # 在用户选择前先显示所有期刊的文件统计信息
    print("正在收集文件统计信息...") 
    
    stats = {}
    for magazine_type, config in MAGAZINE_CONFIG.items():
        try:
            stats[magazine_type] = count_magazine_files(magazine_type, args.output)
        except Exception as e:
            print(f"统计 {config['title']} 信息时出错: {e}")
            stats[magazine_type] = {"total_epub": 0, "unique_dates": 0, "existing_json": 0, "to_process": 0}
    
    # 表格列配置：标题、宽度、对齐方式
    columns = [
        {'title': '期刊名称', 'width': 10, 'align': 'left'},
        {'title': '总期数', 'width': 5, 'align': 'center'},
        {'title': '已生成', 'width': 5, 'align': 'center'},
        {'title': '待处理', 'width': 5, 'align': 'center'}
    ]
    
    def format_cell(text, width, align):
        # 计算显示宽度（中文字符算2个英文字符宽度）
        display_width = sum(2 if '\u4e00' <= c <= '\u9fff' else 1 for c in text)
        padding = width * 2 - display_width
        
        if align == 'left':
            return f"{text}{' ' * padding}"
        else:  # center
            left = padding // 2
            return f"{' ' * left}{text}{' ' * (padding - left)}"
    
    # 计算表格宽度
    header = ' '.join(format_cell(col['title'], col['width'], col['align']) for col in columns)
    divider = '-' * (len(header) + 13)
    
    # 打印统计信息标题和分隔线
    print(f"\n{divider}")
    
    # 打印表头
    print(header)
    print(divider)
    
    # 打印数据行
    for magazine_type, config in MAGAZINE_CONFIG.items():
        s = stats[magazine_type]
        row = [
            config['title'],
            str(s['total_folders']),
            str(s['existing_json']),
            str(s['to_process'])
        ]
        print(' '.join(format_cell(cell, col['width'], col['align']) 
                     for cell, col in zip(row, columns)))
    
    print(divider + "\n")
    
    # 处理杂志选择
    selected_magazine = None
    
    if args.magazine:
        # 如果通过命令行参数指定了杂志
        if args.magazine.isdigit():
            choice = int(args.magazine)
            if 1 <= choice <= len(MAGAZINE_CONFIG):
                selected_magazine = list(MAGAZINE_CONFIG.keys())[choice-1]
            else:
                print(f"错误: 无效的杂志编号 {choice}")
                sys.exit(1)
        elif args.magazine in MAGAZINE_CONFIG:
            selected_magazine = args.magazine
        else:
            print(f"错误: 未知的杂志类型 '{args.magazine}'")
            sys.exit(1)
    else:
        # 交互式选择
        print("请选择要处理的杂志类型:")
        for i, (key, config) in enumerate(MAGAZINE_CONFIG.items(), 1):
            print(f"{i}. {config['name']} ({config['title']})")
        
        while True:
            try:
                choice = input(f"请输入选项编号 (1-{len(MAGAZINE_CONFIG)}), 或按 Ctrl+C 退出: ")
                if choice.isdigit() and 1 <= int(choice) <= len(MAGAZINE_CONFIG):
                    selected_magazine = list(MAGAZINE_CONFIG.keys())[int(choice)-1]
                    break
                print(f"错误: 请输入1到{len(MAGAZINE_CONFIG)}之间的数字")
            except (ValueError, IndexError):
                print("错误: 请输入有效的数字")
            except KeyboardInterrupt:
                print("\n操作已取消")
                sys.exit(0)
    
    # 处理选定的杂志
    if selected_magazine:
        process_magazine(selected_magazine, args.output)

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n程序已被用户中断")
        sys.exit(0)
